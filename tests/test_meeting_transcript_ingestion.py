from pathlib import Path

import app.models  # noqa: F401
from docx import Document as DocxDocument
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db.base import Base
from app.models.chunk import Chunk
from app.schemas.documents import DocumentIngestRequest
from app.schemas.retrieval import RetrievalFilter, SearchRequest
from app.services.ingestion.service import DocumentIngestionService
from app.services.meeting_transcript import enrich_meeting_metadata, extract_meeting_fields, meeting_trace
from app.services.retrieval.service import RetrievalService
from app.services.storage.service import StoredFile


class FakeOpenSearchChunkIndexer:
    def index_chunk(self, chunk, document, version) -> bool:
        return True


def test_meeting_docx_ingestion_preserves_structured_metadata(tmp_path: Path, monkeypatch):
    db_session = _db_session()
    try:
        meeting_path = tmp_path / "项目例会纪要.docx"
        _write_sample_meeting_docx(meeting_path)
        monkeypatch.setattr(
            "app.services.ingestion.service.OpenSearchChunkIndexer",
            lambda: FakeOpenSearchChunkIndexer(),
        )

        job = DocumentIngestionService(db_session).ingest_uploaded_file(
            DocumentIngestRequest(
                source_uri=meeting_path.name,
                title="项目例会纪要",
                source_type="meeting",
                document_type="meeting",
            ),
            StoredFile(
                storage_uri=str(meeting_path),
                local_path=meeting_path,
                file_name=meeting_path.name,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )

        assert job.status == "completed"
        chunks = db_session.query(Chunk).filter(Chunk.document_id == job.document_id).all()
        assert chunks
        meeting_metadata = [chunk.metadata_json for chunk in chunks if chunk.metadata_json.get("meeting_transcript")]
        assert meeting_metadata
        combined_fields = {field for metadata in meeting_metadata for field in metadata["meeting_fields"]}
        assert {"speaker", "timestamp", "decision", "action_item", "risk"} <= combined_fields
        assert all(metadata["transcript_as_fact"] is False for metadata in meeting_metadata)
        assert all(metadata["evidence_required"] is True for metadata in meeting_metadata)
        assert all(metadata.get("source_chunk_id") for metadata in meeting_metadata)
    finally:
        db_session.close()


def test_meeting_retrieval_returns_evidence_metadata(tmp_path: Path, monkeypatch):
    db_session = _db_session()
    try:
        meeting_path = tmp_path / "项目例会纪要.docx"
        _write_sample_meeting_docx(meeting_path)
        monkeypatch.setattr(
            "app.services.ingestion.service.OpenSearchChunkIndexer",
            lambda: FakeOpenSearchChunkIndexer(),
        )
        job = DocumentIngestionService(db_session).ingest_uploaded_file(
            DocumentIngestRequest(
                source_uri=meeting_path.name,
                title="项目例会纪要",
                source_type="meeting",
                document_type="meeting",
            ),
            StoredFile(
                storage_uri=str(meeting_path),
                local_path=meeting_path,
                file_name=meeting_path.name,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )

        response = _fallback_search(
            db_session,
            query="行动项",
            document_id=job.document_id,
            document_type="meeting",
            monkeypatch=monkeypatch,
        )

        assert response.results
        assert response.trace["meeting_transcript_used"] is True
        assert response.trace["action_items_detected"] >= 1
        assert response.trace["transcript_as_fact"] is False
        assert response.results[0].metadata["meeting_transcript"] is True
        assert response.results[0].metadata["source_chunk_id"] == response.results[0].chunk_id
    finally:
        db_session.close()


def test_meeting_extractor_detects_speaker_timestamp_decision_action_and_risk():
    fields = extract_meeting_fields(
        "\n".join(
            [
                "会议时间：2026年5月1日",
                "严总：决定采用数字化交付平台作为试点。",
                "唐总：负责跟进交付标准确认，截止2026年5月10日完成。",
                "风险：客户侧资料不完整，需要待确认。",
            ]
        )
    )

    assert fields["timestamp"] == "2026年5月1日"
    assert "严总" in fields["speakers"]
    assert fields["decision"]
    assert fields["action_item"]
    assert fields["risk"]
    assert "2026年5月10日" in fields["deadline"]


def test_non_meeting_document_is_not_marked_as_meeting():
    metadata = enrich_meeting_metadata(
        text="行动项这个词出现在普通交付标准里，也不能让文档变成会议纪要。",
        metadata={"parser": "docx"},
        source_type="tender",
        document_type="tender",
        source_name="交付标准",
        source_uri="交付标准.docx",
    )

    assert metadata == {"parser": "docx"}


def test_meeting_query_boosts_cover_action_decision_and_risk():
    service = RetrievalService(db=None)

    assert "行动计划" in service._meeting_query_boosts("会议里有哪些行动项？", {"source_type": "meeting"})
    assert "会议结论" in service._meeting_query_boosts("会议里形成了哪些决策？", {"document_type": "meeting"})
    assert "问题一" in service._meeting_query_boosts("会议中提到哪些风险？", {"source_type": "meeting"})
    assert service._meeting_query_boosts("会议中提到哪些风险？", {"source_type": "tender"}) == {}


def test_meeting_trace_never_marks_action_decision_or_risk_as_fact():
    result = type(
        "Result",
        (),
        {
            "chunk_id": "meeting-chunk",
            "metadata": {
                "meeting_transcript": True,
                "meeting_fields_matched": ["action_item", "decision", "risk"],
                "action_item": ["行动项：唐总负责跟进。"],
                "decision": ["决定采用数字化交付平台。"],
                "risk": ["风险：客户侧资料待确认。"],
                "transcript_as_fact": True,
            },
        },
    )()

    trace = meeting_trace([result])

    assert trace["meeting_transcript_used"] is True
    assert trace["action_items_detected"] == 1
    assert trace["decisions_detected"] == 1
    assert trace["risks_detected"] == 1
    assert trace["transcript_as_fact"] is False
    assert trace["evidence_required"] is True


def _fallback_search(db_session, query: str, document_id: str, document_type: str, monkeypatch):
    def fail_sparse(*_args, **_kwargs):
        raise RuntimeError("force database fallback")

    service = RetrievalService(db_session)
    monkeypatch.setattr(service, "_execute_sparse_search", fail_sparse)
    return service.search(
        SearchRequest(
            query=query,
            retrieval_mode="sparse",
            enable_dense=False,
            filters=RetrievalFilter(document_id=document_id, document_type=document_type),
            top_k=3,
        )
    )


def _db_session():
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine)()


def _write_sample_meeting_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_heading("项目例会纪要", level=1)
    document.add_paragraph("会议主题：数字化交付平台试点")
    document.add_paragraph("会议时间：2026年5月1日")
    document.add_paragraph("严总：决定采用数字化交付平台作为试点。")
    document.add_paragraph("行动项：唐总负责跟进交付标准确认，截止2026年5月10日完成。")
    document.add_paragraph("风险：客户侧资料不完整，需要待确认。")
    document.save(path)
