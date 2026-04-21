from pathlib import Path

from docx import Document as DocxDocument

from app.services.parsing.base import ParsedBlock, ParsedDocument


class DocxParser:
    def parse(self, source_path: Path) -> ParsedDocument:
        doc = DocxDocument(source_path)
        blocks: list[ParsedBlock] = []
        heading_path: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                level = self._heading_level(style_name)
                heading_path = heading_path[: max(level - 1, 0)] + [text]
                blocks.append(ParsedBlock(text=text, block_type="heading", title_path=heading_path.copy()))
            else:
                blocks.append(ParsedBlock(text=text, block_type="paragraph", title_path=heading_path.copy()))

        return ParsedDocument(
            source_path=source_path,
            title=source_path.name,
            blocks=blocks,
            metadata={"parser": "docx"},
        )

    def _heading_level(self, style_name: str) -> int:
        suffix = style_name.split()[-1]
        return int(suffix) if suffix.isdigit() else 1

